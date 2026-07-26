# -*- coding: utf-8 -*-
"""Generate Marmara University Tez Önerisi Formu (2026) as DOCX."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def set_run_font(run, size=11, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading_center(text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading_left(text, size=12, bold=True, space_before=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(text, first_indent=True, size=11, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_font(r, size=size)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    return p


def add_bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size=size)
    p.paragraph_format.space_after = Pt(3)
    return p


def shade_cell(cell, color="1F4E79"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    if color:
        r.font.color.rgb = RGBColor(*color)


# =========================
# COVER / HEADER
# =========================
add_heading_center("T.C.", size=12)
add_heading_center("MARMARA ÜNİVERSİTESİ", size=14)
add_heading_center("TEZ ÖNERİSİ FORMU", size=16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Bankacılık Sektörü Müşteri Şikayet Analizi — Metin Madenciliği ve Makine Öğrenmesi (2026)"
)
set_run_font(r, size=10, italic=True)

# =========================
# i. STUDENT INFO
# =========================
add_heading_left("i. ÖĞRENCİ BİLGİLERİ / STUDENT INFORMATION", size=12)

info_rows = [
    ("Ad Soyadı / Name–Surname:", "[Doldurulacak]"),
    ("TC Kimlik Numarası / Identity Number:", "[Doldurulacak]"),
    ("Öğrenci Numarası / Student Number:", "[Doldurulacak]"),
    ("Anabilim Dalı / Department:", "[Doldurulacak]"),
    ("Programı / Program:", "[Doldurulacak]"),
    ("Program Dili / Language of Program:", "Türkçe / Turkish"),
    ("E-Posta / E-Mail:", "[Doldurulacak]"),
    ("Telefon / Phone Number:", "[Doldurulacak]"),
    ("Adres / Address:", "[Doldurulacak]"),
    ("Tez Danışmanı / Supervisor:", "[Doldurulacak]"),
    ("İkinci Tez Danışmanı / Co-Supervisor:", "[Varsa doldurulacak]"),
]
t = doc.add_table(rows=len(info_rows), cols=2)
t.style = "Table Grid"
for i, (k, v) in enumerate(info_rows):
    set_cell_text(t.rows[i].cells[0], k, bold=True, size=10)
    set_cell_text(t.rows[i].cells[1], v, size=10)

# =========================
# ii. THESIS INFO
# =========================
add_heading_left("ii. TEZ BİLGİLERİ / THESIS INFORMATION", size=12)

add_heading_left("Türkçe Tez Başlığı / Title of Thesis in Turkish", size=11, space_before=8)
add_body(
    "Bankacılık Sektöründe Çevrimiçi Müşteri Şikayetlerinin Metin Madenciliği ve Makine Öğrenmesi "
    "Teknikleri ile İncelenmesi: Kuveyt Türk, VakıfBank ve İşBankası Üzerine Karşılaştırmalı Bir Araştırma",
    first_indent=False,
)

add_heading_left("Türkçe Tez Önerisi Özeti", size=11, space_before=8)
tr_ozet = (
    "Bu tez çalışmasının amacı, Türkiye’de faaliyet gösteren üç bankaya (katılım bankası olarak Kuveyt Türk, "
    "kamu bankası olarak VakıfBank ve özel banka olarak İşBankası) yönelik Şikayetvar platformunda yayımlanan "
    "çevrimiçi müşteri şikayetlerini Doğal Dil İşleme (NLP), konu modelleme ve denetimli makine öğrenmesi "
    "teknikleriyle analiz ederek bankalar arası müşteri deneyimi dinamiklerini karşılaştırmalı biçimde "
    "ortaya koymaktır. Çalışma, büyük hacimli Türkçe şikayet metinlerinden otomatik konu çıkarımı yapmak, "
    "şikayetlerin çözülme durumunu tahmin eden sınıflandırma modelleri geliştirmek ve elde edilen bulguları "
    "istatistiksel hipotez testleriyle desteklemek üzere tasarlanmıştır."
)
add_body(tr_ozet)

tr_ozet2 = (
    "Araştırmanın veri kaynağı, Şikayetvar üzerinde 2026 yılına ait banka şikayetleridir. Mevcut proje "
    "veri setinde Ocak–Nisan 2026 döneminde toplanmış yaklaşık 6.396 ham şikayet kaydı bulunmakta; "
    "ön işleme sonrasında analiz için yaklaşık 6.033 temiz kayıt kullanılmaktadır (İşBankası ~3.620, "
    "Kuveyt Türk ~1.657, VakıfBank ~756). Her kayıt; şikayet başlığı ve tam metin, tarih, çözüm durumu "
    "(Çözüldü/Çözülmedi), memnuniyet skoru (1–5), anahtar kelimeler, görüntülenme sayısı, şirket yanıtı "
    "ve URL gibi meta alanları içermektedir. Veriler web scraping ile toplanmış, ardından altı adımlı "
    "analiz boru hattı (kurulum, keşifsel veri analizi, Türkçe metin ön işleme, LDA konu modelleme, "
    "TF-IDF + makine öğrenmesi, karşılaştırmalı hipotez raporlama) ile işlenmiştir."
)
add_body(tr_ozet2)

tr_ozet3 = (
    "Yöntemsel olarak çalışma şu ana hatları izler: (1) kapsamlı Türkçe stop-word listesi, tokenizasyon "
    "ve Snowball kök bulma ile metin temizliği; (2) banka bazında eş anlamlı şikayet kategorilerinin "
    "birleştirilmesi ve keşifsel görselleştirme; (3) Gensim LDA ile banka özelinde optimal konu sayısı "
    "araması (coherence + perplexity) ve bigram tespiti; (4) TF-IDF (unigram+bigram, 10.000 özellik) "
    "ile birlikte meta özellikler (metin uzunluğu, log görüntülenme, şirket yanıtı varlığı) kullanılarak "
    "Lojistik Regresyon, Random Forest ve doğrusal SVM modellerinin eğitimi; (5) Accuracy, Precision, "
    "Recall, F1, ROC-AUC ve 5-kat çapraz doğrulama ile performans değerlendirmesi; (6) H1–H5 hipotezlerinin "
    "Mann-Whitney U, Spearman korelasyon, trend regresyonu ve konu dağılımı karşılaştırmaları ile testi."
)
add_body(tr_ozet3)

tr_ozet4 = (
    "Kuramsal çerçeve; hizmet kalitesi (SERVQUAL), müşteri memnuniyeti (beklenti–algı), dijital itibar "
    "yönetimi ve veri temelli karar verme yaklaşımlarına dayanmaktadır. Çalışmanın özgün değeri, farklı "
    "bankacılık modellerini (katılım–kamu–özel) aynı açık veri kaynağı ve bütünleşik NLP–ML metodolojisi "
    "ile karşılaştırması; betimleyici konu modellemeyi öngörücü çözüm tahmini ile birleştirmesi ve "
    "çarpık memnuniyet dağılımında medyan merkezli istatistik politikası ile metodolojik tutarlılık "
    "sağlamasıdır. Bilimsel katkı olarak Türkçe bankacılık şikayet analizine ampirik bir çerçeve sunulması; "
    "uygulamalı katkı olarak ise bankalara konu bazlı erken uyarı, yanıt performansı ve çözüm tahmini "
    "içgörüleri üretilmesi hedeflenmektedir."
)
add_body(tr_ozet4)

add_heading_left("Türkçe Anahtar Kelimeler / Keywords in Turkish", size=11, space_before=8)
add_body(
    "Metin Madenciliği; Doğal Dil İşleme; Makine Öğrenmesi; Bankacılık Sektörü; Müşteri Şikayetleri; "
    "Konu Modelleme; LDA; TF-IDF; Sınıflandırma; Müşteri Memnuniyeti; Şikayetvar; Karşılaştırmalı Analiz; "
    "Karar Destek; Dijital Bankacılık; Kuveyt Türk; VakıfBank; İşBankası",
    first_indent=False,
)

add_heading_left("İngilizce Tez Başlığı / Title of Thesis in English", size=11, space_before=10)
add_body(
    "Analyzing Online Customer Complaints in the Banking Sector Using Text Mining and Machine Learning Techniques: "
    "A Comparative Study of Kuveyt Türk, VakıfBank and İşBankası",
    first_indent=False,
)

add_heading_left("İngilizce Tez Önerisi Özeti", size=11, space_before=8)
en1 = (
    "The aim of this thesis is to analyze online customer complaints published on the Şikayetvar platform "
    "for three banks operating in Türkiye—Kuveyt Türk (participation bank), VakıfBank (state-owned bank), "
    "and İşBankası (private bank)—by applying Natural Language Processing (NLP), topic modeling, and "
    "supervised machine learning, and to compare customer experience dynamics across institutions. The study "
    "is designed to automatically extract complaint topics from large-scale Turkish text, develop classification "
    "models that predict complaint resolution status, and support findings with statistical hypothesis testing."
)
add_body(en1)

en2 = (
    "The data source consists of bank-related complaints from Şikayetvar for 2026. The current project dataset "
    "covers January–April 2026 and contains approximately 6,396 raw complaint records; after preprocessing, "
    "about 6,033 cleaned records are used for analysis (İşBankası ~3,620; Kuveyt Türk ~1,657; VakıfBank ~756). "
    "Each record includes complaint title and full text, date, resolution status (Resolved/Unresolved), "
    "satisfaction score (1–5), keywords, view count, company reply, and URL. Data are collected via web scraping "
    "and processed through a six-step analytical pipeline: environment setup, exploratory data analysis, "
    "Turkish text preprocessing, LDA topic modeling, TF-IDF with machine learning, and comparative hypothesis reporting."
)
add_body(en2)

en3 = (
    "Methodologically, the thesis follows these core steps: (1) Turkish stop-word filtering, tokenization, "
    "and Snowball stemming; (2) bank-specific synonym merging of complaint categories and exploratory "
    "visualization; (3) Gensim LDA with per-bank optimal topic search using coherence and perplexity, "
    "including bigram detection; (4) TF-IDF representation (unigram+bigram, 10,000 features) combined with "
    "meta-features (token length, log view count, company reply indicator) to train Logistic Regression, "
    "Random Forest, and Linear SVM models with class balancing; (5) evaluation via Accuracy, Precision, "
    "Recall, F1, ROC-AUC, and 5-fold cross-validation; and (6) testing hypotheses H1–H5 using Mann–Whitney U, "
    "Spearman correlation, trend regression, and topic-distribution comparison."
)
add_body(en3)

en4 = (
    "The theoretical framework draws on service quality (SERVQUAL), customer satisfaction (expectation–perception), "
    "digital reputation management, and data-driven decision making. The originality of the study lies in comparing "
    "three distinct banking models (participation–public–private) under a unified open-data NLP–ML methodology; "
    "integrating descriptive topic modeling with predictive resolution classification; and adopting a median-based "
    "central-tendency policy for highly skewed satisfaction scores. Scientifically, the thesis contributes an "
    "empirical framework for Turkish banking complaint analytics; practically, it aims to provide banks with "
    "topic-level early-warning signals, reply-performance insights, and resolution-prediction support for "
    "customer experience management."
)
add_body(en4)

add_heading_left("İngilizce Anahtar Kelimeler / Keywords in English", size=11, space_before=8)
add_body(
    "Text Mining; Natural Language Processing; Machine Learning; Banking Sector; Customer Complaints; "
    "Topic Modeling; LDA; TF-IDF; Classification; Customer Satisfaction; Şikayetvar; Comparative Analysis; "
    "Decision Support; Digital Banking; Kuveyt Türk; VakıfBank; İşBankası",
    first_indent=False,
)

# =========================
# 1. AIM
# =========================
add_heading_left("1. AMAÇ ve HEDEFLER / AIM AND OBJECTIVES", size=12, space_before=14)
add_heading_left("1.1. Araştırmanın Amacı", size=11, space_before=6)
add_body(
    "Bu tez çalışmasının temel amacı, bankacılık sektöründe çevrimiçi platformlarda paylaşılan müşteri "
    "şikayetlerini Doğal Dil İşleme ve Makine Öğrenmesi teknikleri kullanarak analiz etmek ve Kuveyt Türk, "
    "VakıfBank ile İşBankası’nın müşteri deneyimi performanslarını karşılaştırmalı olarak incelemektir. "
    "Araştırma, büyük hacimli Türkçe metin verilerinden otomatik olarak anlamlı bilgi üretmeyi, şikayet "
    "konularını ortaya çıkarmayı ve çözüm başarısını tahmin edebilen modeller geliştirmeyi amaçlamaktadır. "
    "Böylece bankaların müşteri memnuniyetine ilişkin güçlü ve zayıf yönleri veri temelli olarak "
    "değerlendirilebilecektir."
)

add_heading_left("1.2. Araştırmanın Hedefleri", size=11, space_before=6)
add_body(
    "Bu tez kapsamında ulaşılması hedeflenen somut ve ölçülebilir hedefler şunlardır:",
    first_indent=False,
)
hedefler = [
    "2026 yılına ait Şikayetvar şikayet verilerinin toplanması ve analiz edilebilir birleşik veri setinin oluşturulması.",
    "Türkçe’ye özgü kapsamlı stop-word, tokenizasyon ve kök bulma adımlarıyla temiz metin veri yapısının elde edilmesi.",
    "Keşifsel veri analizi ile şikayet hacmi, çözüm oranı, memnuniyet dağılımı, aylık trend ve kategori profillerinin çıkarılması.",
    "Banka bazında LDA konu modelleme ile optimal konu sayısının belirlenmesi ve konu–çözüm ilişkisinin incelenmesi.",
    "TF-IDF + meta özelliklerle Lojistik Regresyon, Random Forest ve SVM modellerinin eğitilerek çözüm durumu tahmininin yapılması.",
    "Modellerin Accuracy, Precision, Recall, F1, ROC-AUC ve çapraz doğrulama metrikleriyle değerlendirilmesi (H3 eşiği: ≥%70 accuracy).",
    "Üç banka arasında konu dağılımı, çözüm performansı, yanıt oranı ve memnuniyet ilişkilerinin istatistiksel karşılaştırılması.",
    "Bulguların yönetici özet panoları ve raporlarla karar destek çıktısına dönüştürülmesi.",
]
for h in hedefler:
    add_bullet(h)

add_heading_left("1.3. Dayanılan Teoriler ve Modeller", size=11, space_before=6)
add_body(
    "Araştırma aşağıdaki kuramsal ve teknik çerçevelere dayanmaktadır:",
    first_indent=False,
)
for titem in [
    "Hizmet Kalitesi Teorisi (SERVQUAL yaklaşımı)",
    "Müşteri Memnuniyeti ve Beklenti–Algı Modeli",
    "İtibar Yönetimi ve Dijital Geri Bildirim Kuramı",
    "Metin Madenciliği ve Doğal Dil İşleme modelleri",
    "Olasılıksal Konu Modelleme (Latent Dirichlet Allocation – LDA)",
    "Denetimli Makine Öğrenmesi modelleri (LR, Random Forest, SVM)",
]:
    add_bullet(titem)

add_heading_left("1.4. Araştırma Hipotezleri", size=11, space_before=6)
add_body("Araştırma kapsamında test edilen temel hipotezler şunlardır:", first_indent=False)
hips = [
    "H1: Şikayet konuları bankalar arasında istatistiksel olarak anlamlı farklılık göstermektedir.",
    "H2: Belirli şikayet kategorileri / çözüm durumu, memnuniyet ve çözülmeme dinamikleri üzerinde anlamlı farklılık yaratmaktadır (Mann–Whitney U).",
    "H3: Makine öğrenmesi modelleri, şikayetlerin çözülme durumunu %70’in üzerinde doğruluk oranıyla tahmin edebilmektedir.",
    "H4: Şikayet yoğunluğu belirli dönemlerde anlamlı artış (veya azalış) trendleri göstermektedir.",
    "H5: Çözüm durumu ile müşteri memnuniyet skoru arasında pozitif yönlü anlamlı ilişki bulunmaktadır (Spearman).",
]
for h in hips:
    add_bullet(h)

add_body(
    "Bu amaç ve hedefler; mevcut proje boru hattında tanımlı çıktılar (temiz veri, LDA modelleri, "
    "sınıflandırıcılar, hipotez raporları ve görseller) üzerinden ölçülebilir, gerçekçi ve tez süresi "
    "içinde ulaşılabilir niteliktedir. Hedef değişken olarak çözüm durumu (is_resolved) seçilmiş; "
    "memnuniyet skoru ise çarpık dağılım nedeniyle ortalama yerine medyan merkezli istatistik politikası "
    "ile değerlendirilmiştir. Bu ayrım, hem tahmin modellerinde veri sızıntısını (leakage) önlemek hem de "
    "betimsel karşılaştırmalarda istatistiksel güvenilirliği artırmak için bilinçli bir yöntemsel tercihtir."
)

# =========================
# 2. LITERATURE
# =========================
add_heading_left(
    "2. KONU, KAPSAM ve LİTERATÜR ÖZETİ / SUBJECT, SCOPE AND LITERATURE SUMMARY",
    size=12,
    space_before=14,
)

add_heading_left("2.1. Konu ve Araştırma Problemi", size=11, space_before=6)
add_body(
    "Bu tez çalışmasının konusu, bankacılık sektöründe çevrimiçi platformlarda paylaşılan müşteri "
    "şikayetlerinin NLP ve makine öğrenmesi teknikleriyle analiz edilmesi ve bankalar arası karşılaştırmalı "
    "değerlendirme yapılmasıdır. Çalışma; Kuveyt Türk, VakıfBank ve İşBankası hakkında Şikayetvar’da "
    "yayımlanan 2026 yılı müşteri şikayetlerini kapsamaktadır."
)
add_body(
    "Araştırma problemi şu şekilde formüle edilmektedir: “Çevrimiçi bankacılık müşteri şikayetleri metin "
    "madenciliği ve makine öğrenmesi teknikleri kullanılarak anlamlı konulara ayrılabilir mi; çözüm başarısı "
    "tahmin edilebilir mi; katılım, kamu ve özel banka modelleri arasında müşteri deneyimi açısından anlamlı "
    "farklılıklar tespit edilebilir mi?” Bu problem, hem veri bilimi hem de hizmet yönetimi literatüründe "
    "güncel ve stratejik bir araştırma alanına karşılık gelmektedir."
)

add_heading_left("2.2. Kapsam ve Sınırlar", size=11, space_before=6)
add_body(
    "Çalışma, Şikayetvar platformunda ilgili üç bankaya ait 2026 yılı şikayetleriyle sınırlıdır. Mevcut "
    "ampirik veri seti Ocak–Nisan 2026 dönemini kapsamakta olup tez süreci içinde dönem genişletilebilir. "
    "Sosyal medya (Twitter/X, Instagram vb.) verileri, bankaların iç çağrı merkezi kayıtları ve operasyonel "
    "finansal veriler kapsam dışındadır. Analiz; şikayet metni ile mevcut meta veriler (çözülme durumu, "
    "memnuniyet skoru, görüntülenme, şirket yanıtı, anahtar kelimeler) üzerinden yürütülür."
)

add_heading_left("2.3. Literatür Özeti ve Kuramsal Arka Plan", size=11, space_before=6)
add_body(
    "Dijitalleşmenin hız kazanmasıyla müşteri deneyimi, bankacılıkta rekabet avantajını belirleyen temel "
    "unsurlardan biri haline gelmiştir. Hizmet kalitesi literatüründe Parasuraman, Zeithaml ve Berry "
    "tarafından geliştirilen SERVQUAL modeli, müşteri beklentileri ile algılanan hizmet arasındaki farkın "
    "memnuniyet üzerindeki etkisini açıklamaktadır. Günümüzde bu yaklaşım, dijital şikayet platformlarındaki "
    "geri bildirimlerin analiz edilmesiyle daha ölçülebilir hale gelmiştir."
)
add_body(
    "Metin madenciliği ve NLP teknikleri, müşteri yorumlarının analizinde yaygınlaşmıştır. TF-IDF, n-gram "
    "ve gömme yöntemleri metinlerin sayısal temsiline olanak tanırken; LDA gibi konu modelleme "
    "algoritmaları büyük metin koleksiyonlarından otomatik tema çıkarımı sağlar (Blei, Ng ve Jordan, 2003). "
    "Denetimli öğrenmede Lojistik Regresyon, Random Forest ve SVM, metin sınıflandırmada güçlü başarımlar "
    "sunmaktadır. Şikayet çözüm tahmini, literatürde complaint resolution prediction ve churn prediction "
    "çalışmalarıyla benzer bir çerçeveye sahiptir; ancak çoğu çalışma şirket içi veriye dayanmakta, açık "
    "kaynaklı çevrimiçi şikayet platformlarını yeterince kullanmamaktadır."
)
add_body(
    "Uluslararası literatürde bankacılık ve e-ticaret sektörlerinde müşteri yorumlarının duygu analizi ve "
    "konu modelleme ile incelendiği çalışmalar bulunmaktadır. Türkiye bağlamında ise özellikle katılım "
    "bankaları ile kamu ve özel bankaların aynı açık veri kaynağı üzerinden büyük ölçekli metin analiziyle "
    "karşılaştırıldığı çalışmalar sınırlıdır. Ayrıca mevcut çalışmaların bir kısmı yalnızca duygu analizi "
    "veya manuel içerik kodlamasıyla sınırlı kalmakta; betimleyici konu çıkarımı ile öngörücü modellemeyi "
    "entegre etmemektedir."
)

add_heading_left("2.4. Literatürdeki Eksiklikler ve Bu Tezin Konumu", size=11, space_before=6)
add_body(
    "Türkiye bağlamında büyük ölçekli açık veri metin analizinin sınırlılığı: Amazon, TripAdvisor vb. "
    "üzerinde yapılan çalışmalar yaygınken, Türkiye’de Şikayetvar gibi platformların bankacılık özelinde "
    "bütünleşik NLP–ML çerçevesiyle değerlendirildiği akademik çalışmalar görece azdır. Bu tez, üç bankayı "
    "kapsayan ve yeniden üretilebilir bir analiz boru hattı sunarak bu boşluğu doldurmayı amaçlar."
)
add_body(
    "Karşılaştırmalı banka analizlerinin yetersizliği: Literatürde çoğu araştırma tek kurum üzerinedir. "
    "Katılım (Kuveyt Türk), kamu (VakıfBank) ve özel (İşBankası) bankacılık modellerinin aynı metodolojiyle "
    "karşılaştırılması, kurumsal yapı–şikayet teması–çözüm performansı ilişkisini görünür kılar. Proje "
    "bulgularında örneğin çözüm oranları ve şirket yanıt oranlarının bankalar arasında belirgin farklılık "
    "göstermesi, bu karşılaştırmanın ampirik değerini desteklemektedir."
)
add_body(
    "İçerik analizi ile tahminlemenin entegre edilmemesi: Birçok çalışma yalnızca LDA ile tema çıkarır; "
    "ancak konu bulgularını çözüm tahmini modelleriyle birleştirmez. Bu tez, konu modelleme çıktılarını "
    "çözüm oranı analizleriyle ilişkilendirir ve TF-IDF tabanlı sınıflandırıcılarla öngörücü boyutu ekler."
)
add_body(
    "Meta verilerin yetersiz kullanımı: Şikayet platformlarında metin dışında çözüm durumu, memnuniyet, "
    "görüntülenme ve kurum yanıtı gibi alanlar bulunur. Bu tez meta özellikleri model girişine eklerken; "
    "memnuniyeti hedef değişken olarak kullanmayarak sızıntı riskini bilinçli biçimde yönetir. Memnuniyet, "
    "hipotez testlerinde (H2/H5) ayrı bir analitik eksende değerlendirilir."
)

add_heading_left("2.5. Güncel Durum ve Önemi", size=11, space_before=6)
add_body(
    "Dijital şikayet platformları, kurum itibarını doğrudan etkileyen kamusal veri kaynaklarıdır. Artan "
    "dijital bankacılık işlem hacmi, müşteri deneyimi sorunlarını daha görünür hale getirmiştir. Bu tez, "
    "sosyal bilimler ile veri bilimi disiplinlerini birleştirerek bankaların müşteri memnuniyeti "
    "stratejilerine veri temelli karar desteği sağlamayı ve disiplinlerarası bir araştırma örneği sunmayı "
    "hedeflemektedir. Ayrıca çalışma, kamuya açık şikayet verilerinin akademik araştırmada etik ve "
    "yeniden üretilebilir biçimde nasıl kullanılabileceğine dair uygulamalı bir örnek oluşturur."
)

# =========================
# 3. ORIGINALITY
# =========================
add_heading_left("3. ÖZGÜN DEĞER / ORIGINALITY OF STUDY", size=12, space_before=14)
add_body(
    "Bu tez çalışmasının özgün değeri; bankacılık sektöründe çevrimiçi müşteri şikayetlerini büyük ölçekli "
    "açık veri üzerinden, bütünleşik bir metin madenciliği ve makine öğrenmesi çerçevesiyle analiz etmesi "
    "ve üç farklı bankacılık modelini karşılaştırmasıdır. YÖK Ulusal Tez Merkezi’nde bankacılıkta müşteri "
    "memnuniyeti ve hizmet kalitesi üzerine çok sayıda tez bulunmakla birlikte, bunların önemli bir kısmı "
    "anket verilerine veya sınırlı örneklemli içerik analizlerine dayanmaktadır. Açık kaynaklı şikayet "
    "platformlarından elde edilen Türkçe metnin, denetimli ve denetimsiz öğrenme teknikleriyle entegre "
    "biçimde ve yeniden üretilebilir bir boru hattı içinde işlendiği çalışmalar görece sınırlıdır."
)
add_body(
    "Tezin özgünlüğü yalnızca veri kaynağından değil, yöntemsel bütünlüğünden de kaynaklanır: (i) banka "
    "özelinde coherence tabanlı LDA konu seçimi ve konu–çözüm ısı haritaları; (ii) TF-IDF + meta özellik "
    "ile çözüm durumu tahmini; (iii) sınıf dengesizliği için class_weight=balanced yaklaşımı; "
    "(iv) çarpık memnuniyet dağılımında medyan politikası; (v) H1–H5 hipotezlerinin tek bir karşılaştırmalı "
    "raporlama adımında birleştirilmesi. Bu yapı, “şikayet konusu nedir?”, “hangi konular daha az "
    "çözülmektedir?” ve “metinden çözüm durumu tahmin edilebilir mi?” sorularına birlikte yanıt üretir."
)
add_body(
    "Kuramsal açıdan çalışma; hizmet kalitesi, müşteri memnuniyeti ve dijital itibar yönetimi yaklaşımlarını "
    "veri bilimi perspektifiyle bütünleştirerek disiplinlerarası bir model önermektedir. Uygulama açısından "
    "ise oluşturulan temiz veri seti, LDA ve sınıflandırma modelleri ile görsel raporlar, bankalar ve "
    "araştırmacılar için tekrar kullanılabilir bir analitik envanter niteliği taşımaktadır."
)

# =========================
# 4. METHOD
# =========================
add_heading_left("4. YÖNTEM / METHODOLOGY", size=12, space_before=14)
add_body(
    "Bu tez, nicel araştırma tasarımına dayalı; veri madenciliği, metin madenciliği ve makine öğrenmesi "
    "tekniklerinin kullanıldığı uygulamalı bir çalışmadır. Hesaplamalı sosyal bilimler yaklaşımı "
    "benimsenmiştir. Araştırma süreci altı ana adımdan oluşur."
)

add_heading_left("4.1. Veri Toplama", size=11, space_before=6)
add_body(
    "Veri kaynağı Şikayetvar’dır. Kuveyt Türk, VakıfBank ve İşBankası şikayetleri web scraping ile "
    "toplanır. Alanlar: id, tarih, kullanıcı adı, başlık, full_text, company_reply, is_resolved, "
    "satisfaction, keywords, view_count, upvote_count, url. Mevcut veri seti 2026 Ocak–Nisan dönemini "
    "kapsar (~6.396 ham kayıt)."
)

add_heading_left("4.2. Keşifsel Veri Analizi (EDA)", size=11, space_before=6)
add_body(
    "Şikayet sayısı, çözüm oranı, memnuniyet dağılımı, aylık trend, şirket yanıt oranı, görüntülenme "
    "dağılımı ve kategori frekansları incelenir. Banka özelinde eş anlamlı kategori birleştirme "
    "(ör. kart markalarının “Kredi Kartı” altında toplanması) uygulanır. Memnuniyet için medyan "
    "merkezi eğilim olarak kullanılır; çözüm–memnuniyet farkı Mann–Whitney U ile test edilir."
)

add_heading_left("4.3. Türkçe Metin Ön İşleme", size=11, space_before=6)
add_body(
    "Küçük harfe dönüştürme, özel karakter/URL temizliği, tokenizasyon, 700+ kelimelik bankacılık ve "
    "platforma özgü Türkçe stop-word çıkarımı, Snowball Turkish stemming ve post-stem filtre uygulanır. "
    "Boş full_text kayıtları description_preview ile kurtarılır. Çıktı: data/processed/veri_temiz.csv."
)

add_heading_left("4.4. LDA Konu Modelleme", size=11, space_before=6)
add_body(
    "Her banka için ayrı LDA modeli kurulur; k=3…9 aralığında coherence (ve perplexity) ile optimal konu "
    "sayısı seçilir. Gensim Phrases ile bigram’lar (“kredi_kartı”, “mobil_uygulama” vb.) tespit edilir. "
    "Konu etiketleri bankacılık domain bilgisiyle yorumlanır; konu–çözüm oranları ve zaman trendleri "
    "raporlanır. Mevcut sonuçlarda optimal k: VakıfBank=6, İşBankası=4, Kuveyt Türk=8."
)

add_heading_left("4.5. TF-IDF ve Denetimli Öğrenme", size=11, space_before=6)
add_body(
    "Hedef değişken is_resolved (Çözüldü=1 / Çözülmedi=0)’dır. Özellikler: TF-IDF (max_features=10.000, "
    "n-gram=(1,2), sublinear_tf) + token_uzunluk + log(view_count+1) + has_reply. Satisfaction kasıtlı "
    "olarak modele eklenmez (leakage riski). Algoritmalar: Lojistik Regresyon, Random Forest, LinearSVC "
    "(olasılık için CalibratedClassifierCV). class_weight='balanced'; %80/%20 stratified split; 5-kat CV. "
    "Metrikler: Accuracy, Precision, Recall, F1, ROC-AUC. H3 eşiği: Accuracy ≥ 0.70."
)

add_heading_left("4.6. Karşılaştırmalı Analiz ve Hipotez Testleri", size=11, space_before=6)
add_body(
    "H1–H5 hipotezleri banka bazında ve karşılaştırmalı olarak test edilir; sonuçlar dashboard ve "
    "özet tablolarla raporlanır. Bu yöntem seti, tezin konu keşfi, tahminleme ve bankalar arası "
    "karşılaştırma hedeflerine doğrudan hizmet eder."
)

# =========================
# 5. IMPACT
# =========================
add_heading_left("5. YAYGIN ETKİ – KATMA DEĞER / EXPECTED IMPACT", size=12, space_before=14)

impact = doc.add_table(rows=3, cols=2)
impact.style = "Table Grid"
set_cell_text(
    impact.rows[0].cells[0],
    "Yaygın Etki Türleri",
    bold=True,
    size=10,
    color=(255, 255, 255),
    center=True,
)
set_cell_text(
    impact.rows[0].cells[1],
    "Tezden Öngörülen/Beklenen Çıktı, Sonuç ve Etkiler",
    bold=True,
    size=10,
    color=(255, 255, 255),
    center=True,
)
shade_cell(impact.rows[0].cells[0], "1F4E79")
shade_cell(impact.rows[0].cells[1], "1F4E79")

set_cell_text(
    impact.rows[1].cells[0],
    "Bilimsel/Akademik Etkiler\n(Makale, Bildiri, Kitap)",
    bold=True,
    size=10,
)
set_cell_text(
    impact.rows[1].cells[1],
    "Türkçe bankacılık şikayet metinlerinde LDA + TF-IDF + denetimli öğrenme bütünleşik çerçevesinin "
    "ampirik uygulamasıyla literatüre katkı beklenmektedir. Bulguların ulusal/uluslararası hakemli "
    "dergilerde makale ve ilgili kongrelerde bildiri olarak sunulması öngörülür. Çalışma; veri bilimi, "
    "yapay zekâ ve işletme/finans alanlarında yeni lisansüstü araştırmalar için metodolojik referans "
    "niteliği taşıyacaktır. Üretilen temiz veri seti, konu modelleri, performans metrikleri ve hipotez "
    "sonuçları yeniden üretilebilir araştırma çıktısı olarak belgelenecektir.",
    size=10,
)

set_cell_text(
    impact.rows[2].cells[0],
    "Uygulamaya Yönelik Etkiler\n(Veri tabanı, karar destek, eğitim vb.)",
    bold=True,
    size=10,
)
set_cell_text(
    impact.rows[2].cells[1],
    "Banka bazında konu dağılımı, çözüm oranı ve yanıt performansı karşılaştırmaları; müşteri deneyimi "
    "ekiplerine erken uyarı ve önceliklendirme içgörüsü sağlar. Çözüm durumu tahmin modelleri, yoğun "
    "şikayet kuyruklarında riskli kayıtların öne alınmasına destek olabilir. Oluşturulan görsel panolar "
    "ve raporlar (executive dashboard) yönetici düzeyinde karar destek çıktısıdır. Temizlenmiş/etiketlenmiş "
    "veri envanteri ve Python analiz boru hattı, üniversite–sektör iş birlikleri, çalıştaylar ve eğitim "
    "programlarında örnek uygulama olarak kullanılabilir. Böylece tez; bilimsel, operasyonel ve eğitimsel "
    "düzeyde katma değer üretme potansiyeline sahiptir.",
    size=10,
)

# =========================
# 6. SCHEDULE
# =========================
add_heading_left("6. İŞ TAKVİMİ / WORK SCHEDULE", size=12, space_before=14)
add_body(
    "Aşağıdaki iş–zaman çizelgesi, tez çalışma süresinin tamamını kapsayacak şekilde 12 aylık bir planı "
    "örneklemektedir. Ay numaraları tez başlangıcına göredir.",
    first_indent=False,
)

schedule_items = [
    ("1", "Literatür taraması ve tez önerisinin netleştirilmesi", "1–2"),
    ("2", "Web scraping aracının geliştirilmesi / güncellenmesi", "2–3"),
    ("3", "Veri toplama (Şikayetvar – 3 banka)", "3–4"),
    ("4", "Veri kalite kontrolü ve birleştirme (kurulum)", "4"),
    ("5", "Keşifsel veri analizi (EDA) ve görselleştirme", "4–5"),
    ("6", "Türkçe metin ön işleme ve stop-word iyileştirme", "5–6"),
    ("7", "LDA konu modelleme ve konu etiketleme", "6–7"),
    ("8", "TF-IDF + ML modelleme ve hiperparametre ayarı", "7–8"),
    ("9", "Hipotez testleri ve karşılaştırmalı analiz", "8–9"),
    ("10", "Bulguların yorumlanması ve sektörel çıkarımlar", "9–10"),
    ("11", "Tez yazımı (yöntem, bulgular, tartışma)", "9–11"),
    ("12", "Danışman geri bildirimleri, düzeltme ve teslim", "11–12"),
]
st = doc.add_table(rows=1 + len(schedule_items), cols=3)
st.style = "Table Grid"
for j, h in enumerate(["İş No", "Yapılacak İşler", "Aylar"]):
    set_cell_text(st.rows[0].cells[j], h, bold=True, size=10, color=(255, 255, 255), center=True)
    shade_cell(st.rows[0].cells[j], "1F4E79")
for i, (no, ish, aylar) in enumerate(schedule_items, start=1):
    set_cell_text(st.rows[i].cells[0], no, size=10, center=True)
    set_cell_text(st.rows[i].cells[1], ish, size=10)
    set_cell_text(st.rows[i].cells[2], aylar, size=10, center=True)

# =========================
# 7. RISK
# =========================
add_heading_left("7. RİSK YÖNETİMİ TABLOSU / RISK MANAGEMENT PLAN", size=12, space_before=14)
risks = [
    (
        "Şikayetvar sayfa yapısı değişimi veya erişim kısıtı nedeniyle veri toplamanın aksaması",
        "B Planı: Daha önce toplanmış arşiv CSV’ler kullanılır; alternatif dönem/örneklem ile analiz sürdürülür; "
        "gerekirse manuel doğrulanmış örneklem genişletilir.",
    ),
    (
        "Sınıf dengesizliği (özellikle İşBankası’nda çözülmeyen şikayetlerin baskınlığı) model performansını düşürebilir",
        "B Planı: class_weight=balanced, alternatif eşikler, Precision–Recall odaklı değerlendirme; gerekirse "
        "SMOTE/undersampling ve ek meta özellikler denenir.",
    ),
    (
        "Türkçe stemming gürültüsü / stop-word yetersizliği konu kalitesini bozabilir",
        "B Planı: Post-stem filtre ve LDA-özel stop listesi güncellenir; coherence düşerse k aralığı ve "
        "min/max belge frekansı yeniden ayarlanır.",
    ),
    (
        "H3 eşiğinin (≥%70 accuracy) tüm banka–algoritma kombinasyonlarında sağlanamaması",
        "B Planı: Banka özelinde en iyi algoritma raporlanır; F1/ROC-AUC ile tamamlayıcı değerlendirme yapılır; "
        "BERTurk vb. transformer tabanlı temsil deneysel ek olarak değerlendirilir.",
    ),
    (
        "Zaman aralığının kısmi olması (ör. yalnızca ilk dört ay) trend genellenebilirliğini sınırlayabilir",
        "B Planı: Tez sürecinde veri dönemi genişletilir; mevcut dönemde banka içi karşılaştırmalar ve "
        "konu–çözüm analizleri öncelenir.",
    ),
]
rt = doc.add_table(rows=1 + len(risks), cols=2)
rt.style = "Table Grid"
set_cell_text(
    rt.rows[0].cells[0],
    "Potansiyel Riskler / Potential Risks",
    bold=True,
    size=10,
    color=(255, 255, 255),
    center=True,
)
set_cell_text(
    rt.rows[0].cells[1],
    "B Planı / Plan-B",
    bold=True,
    size=10,
    color=(255, 255, 255),
    center=True,
)
shade_cell(rt.rows[0].cells[0], "1F4E79")
shade_cell(rt.rows[0].cells[1], "1F4E79")
for i, (risk, plan) in enumerate(risks, start=1):
    set_cell_text(rt.rows[i].cells[0], risk, size=10)
    set_cell_text(rt.rows[i].cells[1], plan, size=10)

# =========================
# 8. ETHICS
# =========================
add_heading_left("8. ETİK KURUL / ETHIC COMMITTEE", size=12, space_before=14)
add_body(
    "☐ Bu çalışma için Etik Kurul onayına gerek vardır.\n"
    "☑ Bu çalışma için Etik Kurulu onayına gerek yoktur.",
    first_indent=False,
)
add_body(
    "Gerekçe: Araştırma, kamuya açık Şikayetvar platformunda yayımlanmış şikayet metinleri ve meta "
    "veriler üzerinde yürütülmektedir; bireylere yönelik anket, görüşme veya deneysel müdahale "
    "içermemektedir. Yine de danışman yönlendirmesi doğrultusunda Enstitü/Üniversite etik kurallarına "
    "uyulacak; gerekli görülmesi halinde Etik Kurul başvurusu yapılacaktır. Kişisel tanımlayıcı "
    "bilgilerin raporlarda anonimleştirilmesine özen gösterilecektir."
)

# =========================
# 9. REFERENCES
# =========================
add_heading_left("9. KAYNAKLAR / REFERENCES", size=12, space_before=14)
add_body(
    "(Kaynak gösterimi, Marmara Üniversitesi ilgili Enstitü/Fakülte Tez Yazım Kılavuzu formatına göre "
    "nihai tezde güncellenecektir. Aşağıda örnek temel kaynaklar listelenmiştir.)",
    first_indent=False,
    size=10,
)

refs = [
    "Blei, D. M., Ng, A. Y., & Jordan, M. I. (2003). Latent Dirichlet allocation. Journal of Machine Learning Research, 3, 993–1022.",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
    "Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20(3), 273–297.",
    "Parasuraman, A., Zeithaml, V. A., & Berry, L. L. (1988). SERVQUAL: A multiple-item scale for measuring consumer perceptions of service quality. Journal of Retailing, 64(1), 12–40.",
    "Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text retrieval. Information Processing & Management, 24(5), 513–523.",
    "Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to information retrieval. Cambridge University Press.",
    "Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
    "Řehůřek, R., & Sojka, P. (2010). Software framework for topic modelling with large corpora. LREC 2010 Workshop on New Challenges for NLP Frameworks.",
    "Oliver, R. L. (1980). A cognitive model of the antecedents and consequences of satisfaction decisions. Journal of Marketing Research, 17(4), 460–469.",
    "Agresti, A. (2013). Categorical data analysis (3rd ed.). Wiley.",
    "Jurafsky, D., & Martin, J. H. (2023). Speech and language processing (3rd ed. draft). Stanford University.",
    "Şikayetvar. (2026). Çevrimiçi müşteri şikayet platformu. https://www.sikayetvar.com",
]
for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    set_run_font(r, size=10)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(4)

# =========================
# PROJECT HIGH-LEVEL APPENDIX
# =========================
add_heading_left(
    "EK A — PROJE YÜKSEK DÜZEY ÖZETİ (UYGULAMA BORU HATTI)",
    size=12,
    space_before=16,
)
add_body(
    "Bu ek, tez önerisine temel oluşturan mevcut yazılım projesinin yüksek düzey mimarisini özetler.",
    first_indent=False,
)

add_heading_left("A.1. Proje Bileşenleri", size=11, space_before=6)
comps = [
    "sikayetvar-web-scraper/: Şikayetvar’dan banka bazlı ham CSV toplama scriptleri",
    "adim1_kurulum.py: klasör yapısı, CSV normalizasyonu, NLTK, stop-word üretimi",
    "adim2_eda.py: keşifsel analiz ve figürler (01–09)",
    "adim3_onisleme.py: Türkçe NLP ön işleme ve kelime bulutu/n-gram çıktıları (10–16)",
    "adim4_lda.py: banka bazlı LDA, coherence, pyLDAvis, konu–çözüm analizi (17–21)",
    "adim5_tfidf.py: TF-IDF + LR/RF/SVM, ROC/karışıklık matrisi/özellik önemi (23–29)",
    "adim6_karsilastirma.py: H1–H5 hipotez testleri ve yönetici panoları (30–39)",
    "config.py / requirements.txt: merkezi parametreler ve bağımlılıklar",
]
for c in comps:
    add_bullet(c)

add_heading_left("A.2. Mevcut Veri Özeti (2026)", size=11, space_before=6)
dt = doc.add_table(rows=4, cols=5)
dt.style = "Table Grid"
headers = ["Banka", "Ham Kayıt", "Temiz Kayıt*", "Çözüm Oranı*", "Yanıt Oranı*"]
for j, h in enumerate(headers):
    set_cell_text(dt.rows[0].cells[j], h, bold=True, size=9, color=(255, 255, 255), center=True)
    shade_cell(dt.rows[0].cells[j], "1F4E79")
rows_data = [
    ("VakıfBank", "809", "756", "%54,5", "%100,0"),
    ("İşBankası", "3.781", "3.620", "%22,6", "%17,1"),
    ("Kuveyt Türk", "1.806", "1.657", "%42,3", "%86,9"),
]
for i, row in enumerate(rows_data, start=1):
    for j, val in enumerate(row):
        set_cell_text(dt.rows[i].cells[j], val, size=9, center=True)
add_body(
    "* Temiz kayıt ve oranlar processed veri / genel özet metriklerine göredir. Dönem: 2026-01-01 … 2026-04-30.",
    first_indent=False,
    size=9,
)

add_heading_left("A.3. Hipotez Sonuçlarının Özeti (Mevcut Çalıştırma)", size=11, space_before=6)
ht = doc.add_table(rows=6, cols=3)
ht.style = "Table Grid"
for j, h in enumerate(["Hipotez", "Sonuç", "Kısa Kanıt"]):
    set_cell_text(ht.rows[0].cells[j], h, bold=True, size=9, color=(255, 255, 255), center=True)
    shade_cell(ht.rows[0].cells[j], "1F4E79")
hrows = [
    ("H1", "DESTEKLENDİ", "Optimal k farklı: VB=6, İB=4, KT=8"),
    ("H2", "DESTEKLENDİ", "Mann–Whitney p<0.001; çözülen medyan memnuniyet daha yüksek"),
    ("H3", "KISMEN", "9 modelden bir kısmı ≥%70 accuracy eşiğini aşıyor"),
    ("H4", "KISMEN", "İşBankası ve Kuveyt Türk’te anlamlı trend; VakıfBank’ta anlamsız"),
    ("H5", "DESTEKLENDİ", "Spearman ρ≈0,53–0,56; p<0.001 tüm bankalarda"),
]
for i, row in enumerate(hrows, start=1):
    for j, val in enumerate(row):
        set_cell_text(ht.rows[i].cells[j], val, size=9)

add_heading_left("A.4. Teknoloji Yığını", size=11, space_before=6)
add_body(
    "Python; pandas, numpy, scikit-learn, scipy; gensim, pyLDAvis; snowballstemmer, NLTK; "
    "matplotlib, seaborn, wordcloud, plotly; joblib; (opsiyonel ileri aşama) transformers/BERTurk, torch.",
    first_indent=False,
)

# =========================
# 10. APPROVAL
# =========================
add_heading_left("10. ONAY / CONFIRMATION", size=12, space_before=16)
ot = doc.add_table(rows=4, cols=4)
ot.style = "Table Grid"
for j, h in enumerate(["", "Unvan Ad SOYAD", "Tarih", "İmza"]):
    set_cell_text(ot.rows[0].cells[j], h, bold=True, size=10, color=(255, 255, 255), center=True)
    shade_cell(ot.rows[0].cells[j], "1F4E79")
for i, label in enumerate(
    ["Danışman / Supervisor", "İkinci Tez Danışmanı / Co-Supervisor", "Öğrenci / Student"],
    start=1,
):
    set_cell_text(ot.rows[i].cells[0], label, bold=True, size=9)
    for j in range(1, 4):
        set_cell_text(ot.rows[i].cells[j], "", size=9)

out = Path("Tez_Öneri_Banks_2026.docx")
doc.save(str(out))
print("Saved", out.resolve(), "size", out.stat().st_size)


def wc(s):
    return len(re.findall(r"\S+", s))


print("TR ozet words ~", wc(tr_ozet + tr_ozet2 + tr_ozet3 + tr_ozet4))
print("EN ozet words ~", wc(en1 + en2 + en3 + en4))
